"""
Exportación de resúmenes a PDF y DOCX.

Flujo: Markdown → parse → lista de bloques tipados → render a formato destino.
"""

import re
from dataclasses import dataclass, field
from io import BytesIO
from typing import List, Tuple

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from fpdf import FPDF

# ─── Markdown parser ──────────────────────────────────────────────────────────

HEADING_RE = re.compile(r"^(#{2,4})\s+(.+)$")
UL_RE = re.compile(r"^- ")
OL_RE = re.compile(r"^\d+\.\s")
BQ_RE = re.compile(r"^>\s?")
CONTINUATION_RE = re.compile(r"^\s{2,}")

# Inline: code, bold, italic (no anidados)
INLINE_RE = re.compile(r"`([^`\n]+)`|\*\*([^*\n]+)\*\*|\*([^*\n]+)\*")


@dataclass
class MdBlock:
    """Bloque tipado del parser."""
    kind: str  # 'h2' | 'h3' | 'h4' | 'p' | 'ul' | 'ol' | 'blockquote'
    text: str = ""  # para h/p/blockquote
    items: List[str] = field(default_factory=list)  # para ul/ol


def parse_markdown(md: str) -> List[MdBlock]:
    """Convierte Markdown en una lista de MdBlock."""
    lines = md.split("\n")
    blocks: List[MdBlock] = []
    i = 0

    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue

        # Headings
        m = HEADING_RE.match(line)
        if m:
            lvl = len(m.group(1))
            blocks.append(MdBlock(kind=f"h{lvl}", text=m.group(2)))
            i += 1
            continue

        # Blockquote
        if BQ_RE.match(line):
            bq_lines: list[str] = []
            while i < len(lines) and BQ_RE.match(lines[i].strip()):
                bq_lines.append(BQ_RE.sub("", lines[i].strip()))
                i += 1
            blocks.append(MdBlock(kind="blockquote", text=" ".join(bq_lines)))
            continue

        # Unordered list
        if UL_RE.match(line):
            items: list[str] = []
            while i < len(lines) and UL_RE.match(lines[i].strip()):
                item_text = lines[i].strip()[2:]
                i += 1
                while (
                    i < len(lines)
                    and lines[i].strip()
                    and not UL_RE.match(lines[i].strip())
                    and not OL_RE.match(lines[i].strip())
                    and not HEADING_RE.match(lines[i].strip())
                    and not BQ_RE.match(lines[i].strip())
                    and CONTINUATION_RE.match(lines[i])
                ):
                    item_text += " " + lines[i].strip()
                    i += 1
                items.append(item_text)
            blocks.append(MdBlock(kind="ul", items=items))
            continue

        # Ordered list
        if OL_RE.match(line):
            items = []
            while i < len(lines) and OL_RE.match(lines[i].strip()):
                item_text = OL_RE.sub("", lines[i].strip(), count=1)
                i += 1
                while (
                    i < len(lines)
                    and lines[i].strip()
                    and not OL_RE.match(lines[i].strip())
                    and not UL_RE.match(lines[i].strip())
                    and not HEADING_RE.match(lines[i].strip())
                    and not BQ_RE.match(lines[i].strip())
                    and CONTINUATION_RE.match(lines[i])
                ):
                    item_text += " " + lines[i].strip()
                    i += 1
                items.append(item_text)
            blocks.append(MdBlock(kind="ol", items=items))
            continue

        # Paragraph
        p_lines: list[str] = []
        while i < len(lines):
            t = lines[i].strip()
            if not t:
                break
            if (
                HEADING_RE.match(t)
                or UL_RE.match(t)
                or OL_RE.match(t)
                or BQ_RE.match(t)
            ):
                break
            p_lines.append(t)
            i += 1
        if p_lines:
            blocks.append(MdBlock(kind="p", text=" ".join(p_lines)))

    return blocks


def strip_inline(text: str) -> str:
    """Elimina marcadores inline de Markdown (**, *, `) dejando solo el texto."""
    text = re.sub(r"`([^`\n]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*\n]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*\n]+)\*", r"\1", text)
    return text


def parse_inline(text: str) -> List[Tuple[str, str]]:
    """Tokeniza inline markdown en segmentos (tipo, texto).

    Tipos: 'text', 'bold', 'italic', 'code'.
    """
    segments: List[Tuple[str, str]] = []
    last = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > last:
            segments.append(("text", text[last : m.start()]))
        if m.group(1) is not None:
            segments.append(("code", m.group(1)))
        elif m.group(2) is not None:
            segments.append(("bold", m.group(2)))
        elif m.group(3) is not None:
            segments.append(("italic", m.group(3)))
        last = m.end()
    if last < len(text):
        segments.append(("text", text[last:]))
    return segments


# ─── PDF renderer (fpdf2) ─────────────────────────────────────────────────────

FONT_DIR = "/usr/share/fonts/truetype/dejavu"


def _find_font_dir() -> str:
    """Busca el directorio de fuentes DejaVu. Fallback a fuentes del sistema."""
    import os

    candidates = [
        "/usr/share/fonts/truetype/dejavu",
        "/usr/share/fonts/truetype/DejaVu",
        "/usr/local/share/fonts/dejavu",
        "/usr/local/texlive/2025/texmf-dist/fonts/truetype/public/dejavu",
    ]
    for d in candidates:
        if os.path.isdir(d):
            return d
    raise FileNotFoundError(
        "DejaVu fonts not found. Install fonts-dejavu-core."
    )


def render_pdf(blocks: List[MdBlock], title: str) -> bytes:
    """Genera un PDF a partir de bloques de Markdown."""
    font_dir = _find_font_dir()
    pdf = FPDF()
    pdf.add_font("dejavu", "", f"{font_dir}/DejaVuSans.ttf")
    pdf.add_font("dejavu", "B", f"{font_dir}/DejaVuSans-Bold.ttf")
    pdf.add_font("dejavu", "I", f"{font_dir}/DejaVuSans-Oblique.ttf")
    pdf.add_font("dejavu", "BI", f"{font_dir}/DejaVuSans-BoldOblique.ttf")
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()

    # Título del resumen
    pdf.set_font("dejavu", "B", 16)
    pdf.multi_cell(0, 8, title)
    pdf.ln(6)

    for block in blocks:
        if block.kind in ("h2", "h3", "h4"):
            sizes = {"h2": 14, "h3": 12, "h4": 11}
            spacings = {"h2": (6, 3), "h3": (5, 2), "h4": (4, 2)}
            before, after = spacings[block.kind]
            pdf.ln(before)
            pdf.set_font("dejavu", "B", sizes[block.kind])
            pdf.multi_cell(0, 6, strip_inline(block.text))
            pdf.ln(after)

        elif block.kind == "p":
            pdf.set_font("dejavu", "", 10)
            _write_inline_pdf(pdf, block.text)
            pdf.ln(4)

        elif block.kind == "blockquote":
            saved_x = pdf.get_x()
            saved_margin = pdf.l_margin
            pdf.set_left_margin(saved_margin + 10)
            pdf.set_x(saved_margin + 10)
            pdf.set_font("dejavu", "I", 10)
            _write_inline_pdf(pdf, block.text)
            pdf.set_left_margin(saved_margin)
            pdf.ln(4)

        elif block.kind in ("ul", "ol"):
            pdf.set_font("dejavu", "", 10)
            for idx, item in enumerate(block.items):
                prefix = f"{idx + 1}. " if block.kind == "ol" else "• "
                saved_margin = pdf.l_margin
                pdf.set_left_margin(saved_margin + 8)
                pdf.set_x(saved_margin + 8)
                # Escribir prefijo en texto plano y luego el contenido
                pdf.set_font("dejavu", "", 10)
                _write_inline_pdf(pdf, prefix + item)
                pdf.set_left_margin(saved_margin)
                pdf.ln(1)
            pdf.ln(2)

    return bytes(pdf.output())


def _write_inline_pdf(pdf: FPDF, text: str):
    """Escribe texto con formato inline (bold, italic, code) en el PDF."""
    segments = parse_inline(text)
    for seg_type, seg_text in segments:
        if seg_type == "bold":
            pdf.set_font("dejavu", "B", 10)
        elif seg_type == "italic":
            pdf.set_font("dejavu", "I", 10)
        elif seg_type == "code":
            pdf.set_font("dejavu", "", 9)
        else:
            pdf.set_font("dejavu", "", 10)
        pdf.write(5, seg_text)
    pdf.ln()


# ─── DOCX renderer (python-docx) ──────────────────────────────────────────────


def render_docx(blocks: List[MdBlock], title: str) -> bytes:
    """Genera un DOCX a partir de bloques de Markdown."""
    doc = DocxDocument()

    # Ajustar estilo Normal
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Título
    doc.add_heading(title, level=0)

    for block in blocks:
        if block.kind in ("h2", "h3", "h4"):
            level = int(block.kind[1])
            doc.add_heading(strip_inline(block.text), level=level)

        elif block.kind == "p":
            p = doc.add_paragraph()
            _add_inline_runs(p, block.text)

        elif block.kind == "blockquote":
            # Usar estilo Quote si existe, si no Intense Quote, si no párrafo con cursiva
            try:
                p = doc.add_paragraph(style="Quote")
            except KeyError:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Pt(36)
            _add_inline_runs(p, block.text, default_italic=True)

        elif block.kind == "ul":
            for item in block.items:
                p = doc.add_paragraph(style="List Bullet")
                _add_inline_runs(p, item)

        elif block.kind == "ol":
            for item in block.items:
                p = doc.add_paragraph(style="List Number")
                _add_inline_runs(p, item)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_inline_runs(paragraph, text: str, default_italic: bool = False):
    """Añade runs con formato inline a un párrafo DOCX."""
    segments = parse_inline(text)
    for seg_type, seg_text in segments:
        run = paragraph.add_run(seg_text)
        if seg_type == "bold":
            run.bold = True
        elif seg_type == "italic" or (seg_type == "text" and default_italic):
            run.italic = True
        elif seg_type == "code":
            run.font.name = "Courier New"
            run.font.size = Pt(9)